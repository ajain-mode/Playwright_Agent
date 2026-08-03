"""
Generate complete QA AI Agent handover Word document for Mode Global QA team.
Output: docs/handover/QA-AI-Agent-Handover.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
HANDOVER_DIR = ROOT / "handover"
OUT_PATH = HANDOVER_DIR / "QA-AI-Agent-Handover.docx"
DIAGRAM = ROOT / "presentations" / "agentic-pipeline-diagram.png"
if not DIAGRAM.exists():
    DIAGRAM = ROOT / "presentations" / "pptx-v2-captures" / "07-diagram.png"

NAVY = RGBColor(0x1A, 0x36, 0x5D)


def set_heading(p, level: int = 1) -> None:
    p.style = f"Heading {level}"


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Title page ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("QA AI Agent — Handover Document\n")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Playwright Agentic Test Automation for SunTeck TMS\n").font.size = Pt(14)
    sub.add_run("Mode Global · Proof of Concept\n\n").font.size = Pt(12)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: R Systems International\n").font.size = Pt(11)
    meta.add_run(f"Prepared for: Mode Global QA Team\n").font.size = Pt(11)
    meta.add_run(f"Date: {date.today().strftime('%B %d, %Y')}\n").font.size = Pt(11)
    meta.add_run("Classification: Confidential\n").font.size = Pt(11)
    doc.add_page_break()

    # ── 1. Purpose ──
    doc.add_heading("1. Purpose of This Document", level=1)
    doc.add_paragraph(
        "This document is the complete handover guide for the QA AI Agent developed during the Mode Global "
        "proof of concept (POC). It explains how the agentic system was built, how each component works, "
        "what was delivered, and how Mode Global and partner QA teams can install, configure, and use the "
        "agent to generate Playwright test specification (.spec.ts) files from detailed test cases going forward."
    )
    doc.add_paragraph(
        "POC outcomes: 28 executable test scripts successfully generated — 24 billing toggle and 4 DFB — "
        "using the hybrid rule-based + LLM multi-agent pipeline described in this document."
    )

    # ── 2. Prerequisites ──
    doc.add_heading("2. Prerequisites & Environment Setup", level=1)

    doc.add_heading("2.1 Required Software", level=2)
    add_bullet(doc, "Node.js 18+ and npm (project uses TypeScript + Playwright)")
    add_bullet(doc, "Google Chrome (Playwright runs via channel: 'chrome')")
    add_bullet(doc, "Git")
    add_bullet(doc, "Claude CLI (claude command) — primary LLM engine for script generation")
    add_bullet(doc, "Optional: Anthropic API key (ANTHROPIC_API_KEY) — secondary fallback if CLI unavailable")
    add_bullet(doc, "Optional: Cursor IDE — governance rules in .cursor/rules/ guide agent behaviour")

    doc.add_heading("2.2 Clone Repository & Install Dependencies", level=2)
    add_code(doc, "git clone <repository-url>")
    add_code(doc, "cd Playwright_Agent")
    add_code(doc, "npm install")
    add_code(doc, "npx playwright install chrome")

    doc.add_heading("2.3 Git Configuration (One-Time per Clone)", level=2)
    add_code(doc, "git config core.ignorecase false")
    add_code(doc, 'git update-index --skip-worktree src/testmo/testmo-reporter.ts')

    doc.add_heading("2.4 Environment Variables (.env)", level=2)
    doc.add_paragraph(
        "Create a .env file in the project root (this file is gitignored). Required variables for test execution:"
    )
    add_table(doc,
              ["Variable", "Purpose"],
              [
                  ["GLOBAL_PASSWORD", "BTMS global user password"],
                  ["BTMS_SSO_PASSWORD", "BTMS SSO login password"],
                  ["TNX_PASSWORD", "Tender Exchange password"],
                  ["TNX_REP_PASSWORD", "TNX rep password"],
                  ["DME_PASSWORD", "DME application password"],
                  ["TRITAN_CUSTOMER_PASSWORD", "Tritan customer password"],
                  ["ANTHROPIC_API_KEY", "Optional — Anthropic API fallback for LLM generation"],
              ])
    doc.add_paragraph(
        "User credentials (usernames, non-sensitive config) are in src/loginHelpers/userConfig.json (gitignored). "
        "Contact your team lead for the credential files."
    )

    doc.add_heading("2.5 Claude CLI Installation", level=2)
    doc.add_paragraph(
        "The AI Script Writer uses Claude for four strategic generation calls. At startup, the agent runs "
        "claude --version to verify the CLI is available. If not found, generation falls back to rule-based mode only."
    )
    add_numbered(doc, "Install Claude Code / Claude CLI from Anthropic (https://docs.anthropic.com/en/docs/claude-code)")
    add_numbered(doc, "Sign in with your Anthropic account (Max plan or API access)")
    add_numbered(doc, "Verify installation:")
    add_code(doc, "claude --version")
    add_numbered(doc, "The agent invokes Claude as: claude -p --output-format text --tools \"\" --model claude-sonnet-4-6")
    add_numbered(doc, "If Claude CLI is unavailable, set ANTHROPIC_API_KEY in .env — the agent tries the Anthropic Messages API first, then falls back to CLI")

    doc.add_heading("2.6 Verify Agent Setup", level=2)
    add_code(doc, "npm run agent:cli")
    doc.add_paragraph("At startup you should see either:")
    add_bullet(doc, "LLM Service initialized via Claude CLI (version …) — LLM tiers enabled")
    add_bullet(doc, "LLM Service disabled — rule-based generation only (install Claude CLI or set ANTHROPIC_API_KEY)")

    doc.add_page_break()

    # ── 3. What Was Built ──
    doc.add_heading("3. What Was Built in the POC", level=1)
    doc.add_paragraph(
        "R Systems developed a hybrid agentic test automation pipeline inside the existing SunTeck TMS Playwright "
        "framework. Written test cases (from Testmo, CSV, Excel, JSON, or plain text) flow through a multi-agent "
        "pipeline and emerge as validated, runnable Playwright .spec.ts scripts."
    )
    add_bullet(doc, "Four specialised validation/mapping agents + AI Script Writer (LLMService)")
    add_bullet(doc, "130+ declarative step-to-code automation patterns (StepMappings)")
    add_bullet(doc, "5,900+ real UI element IDs indexed from BTMS + DME application source repos")
    add_bullet(doc, "25+ auto-fix quality rules (SpecValidator sanitizer + guardrails)")
    add_bullet(doc, "Cursor governance rules enforcing architecture standards on every generation")
    add_bullet(doc, "12 product test categories configured (DFB, billing toggle, EDI, commission, etc.)")
    add_bullet(doc, "28 POC scripts delivered: 24 billing toggle + 4 DFB in src/tests/AIAgent/")
    add_bullet(doc, "Per-category runtime test data CSVs in src/data/<category>/")
    add_bullet(doc, "Canonical testcase catalog: src/agent/examples/sample-testcases.csv")

    # ── 4. Architecture ──
    doc.add_heading("4. Agentic Architecture", level=1)
    doc.add_paragraph(
        "Generation is orchestrated by PlaywrightAgent (src/agent/PlaywrightAgent.ts). "
        "The pipeline combines deterministic rule-based generation with Claude LLM calls at strategic points. "
        "All output is validated before saving."
    )

    if DIAGRAM.exists():
        doc.add_paragraph("End-to-end pipeline diagram:")
        doc.add_picture(str(DIAGRAM), width=Inches(6.5))
        doc.add_paragraph()

    doc.add_heading("4.1 High-Level Flow", level=2)
    add_numbered(doc, "Test Case Reader (TestCaseParser) — parses input, detects category, extracts fields")
    add_numbered(doc, "App Context Sync (RepoCloneManager) — clones/updates BTMS (mono.git) + DME (dme.git) repos")
    add_numbered(doc, "UI Element Index (AppSourceIndexer) — scans PHP/Twig for real HTML element IDs")
    add_numbered(doc, "CSV Data Service — ensures runtime test data row exists in category CSV")
    add_numbered(doc, "Script Assembler (CodeGenerator) — builds draft .spec.ts via tiers below")
    add_numbered(doc, "③ Quality Guardian (SpecValidator) — validates, auto-fixes, up to 2 correction passes")
    add_numbered(doc, "④ Coverage Ensurer (SpecFeedbackLoop) — compares script vs testcase, fills up to 8 missing steps")
    add_numbered(doc, "Output — src/tests/AIAgent/<category>/<TEST_ID>.spec.ts")

    doc.add_heading("4.2 Generation Tiers (CodeGenerator)", level=2)
    add_table(doc,
              ["Tier", "Method", "When Used"],
              [
                  ["Tier 1", "Full-spec LLM (Call 2)", "First attempt when ≤40 steps and reference spec exists"],
                  ["Tier 2", "Clone + adapt reference (Call 3)", "When new case ≥70% similar to existing automated test"],
                  ["Tier 2b", "Rule-based step-by-step", "Step Interpreter + Code Mapper + 130+ patterns"],
                  ["Tier 3", "Per-step LLM (Call 4)", "Last resort when no pattern matches a step"],
              ])

    doc.add_heading("4.3 What Each Agent Does", level=2)

    doc.add_heading("✦ AI Script Writer (LLMService · LLMPrompts)", level=3)
    doc.add_paragraph(
        "Invokes Claude at four strategic points. Uses governed prompts with framework schema, quality rules, "
        "and reference test patterns. Claude CLI or Anthropic API. If a call fails, the pipeline falls back automatically."
    )
    add_table(doc,
              ["Call", "Business Name", "What It Does", "When It Runs"],
              [
                  ["Call 1", "Data Enricher", "Fills missing test data fields not parsed from testcase", "Before generation when data incomplete"],
                  ["Call 2", "Full Script Writer", "Writes entire script using reference as template", "First attempt (≤40 steps, reference exists)"],
                  ["Call 3", "Smart Adapter", "Clones reference; rewrites only differing steps", "When ≥70% similar to existing test"],
                  ["Call 4", "Step Writer", "Writes code for individual unmatched steps", "Last resort after rules fail"],
              ])
    doc.add_paragraph(
        "Generation order: Call 2 → Call 3 → Rule-based (Step Interpreter + Code Mapper) → Call 4. "
        "Safety net: All AI output passes Quality Guardian and Coverage Ensurer before saving."
    )

    doc.add_heading("① Step Interpreter (Agent 1 · StepProcessor)", level=3)
    doc.add_paragraph(
        "Reads each raw test step and classifies intent. Tracks immutable context snapshots (contextBefore / contextAfter): "
        "current page, edit vs view mode, current tab, current application (BTMS/DME/TNX). "
        "Output: structured ProcessedStep[] consumed by Code Mapper and Quality Guardian."
    )

    doc.add_heading("② Code Mapper (Agent 2 · POMMethodMatcher)", level=3)
    doc.add_paragraph(
        "Maps understood steps to proven Page Object Model methods using StepMappings (130+ patterns), "
        "AppSourceIndexer data (5,900+ UI elements), and reference specs. "
        "Proposes new POM methods with real locators from application source when needed. "
        "Receives correction requests from Quality Guardian and Coverage Ensurer (≤2 iterations via correctBatch())."
    )

    doc.add_heading("③ Quality Guardian (Agent 3 · SpecValidator)", level=3)
    doc.add_paragraph(
        "Validates generated script against guardrails. Runs 13 sanitizer rules (SAN-001–SAN-013) then structural validation. "
        "Auto-fixes violations (navigation, hardcoded values, locator-in-spec issues, billing toggle user switch, etc.). "
        "Inputs: draft script + ProcessedStep[] + POMMethodMatcher. Up to 2 correction passes. "
        "Loops back to Code Mapper on violations."
    )

    doc.add_heading("④ Coverage Ensurer (Agent 4 · SpecFeedbackLoop)", level=3)
    doc.add_paragraph(
        "Compares finished script against original testcase steps. Injects code for missed steps (max 8 injections, 1 pass). "
        "Flags uncertain matches for human review. If changes applied, re-runs Quality Guardian (1 pass). "
        "Ensures every explicit expected result from the testcase has an assertion."
    )

    doc.add_heading("4.4 Self-Healing & Governance", level=2)
    add_bullet(doc, "SpecValidator SAN-001–SAN-013: sanitizer pre-pass (quotes, imports, navigation patterns, etc.)")
    add_bullet(doc, "25+ guardrail rules: no locators in specs, no hardcoded assertion values, POM JSDoc required, etc.")
    add_bullet(doc, "Cursor rules (.cursor/rules/agent-architecture.mdc, CLAUDE.md): architecture standards always applied")
    add_bullet(doc, "Human-authored POM methods are protected — agent may only add new @author AI Agent methods")

    doc.add_page_break()

    # ── 5. Repository Structure ──
    doc.add_heading("5. Repository Structure (Key Paths)", level=1)
    add_table(doc,
              ["Path", "Purpose"],
              [
                  ["src/agent/", "Agent pipeline — parsers, generators, validators, CLI"],
                  ["src/agent/run-agent.ts", "CLI entry point for all agent modes"],
                  ["src/agent/PlaywrightAgent.ts", "Main orchestration class"],
                  ["src/agent/config/StepMappings.ts", "130+ step-to-code patterns"],
                  ["src/agent/config/FieldRegistry.ts", "Unified field name mappings"],
                  ["src/agent/examples/sample-testcases.csv", "Canonical testcase step catalog"],
                  ["src/tests/AIAgent/", "Generated Playwright specs (28 POC scripts)"],
                  ["src/data/<category>/", "Runtime test data CSVs per module"],
                  ["src/pages/", "Page Object Model classes"],
                  ["src/utils/globalConstants.ts", "Assertion constants (never hardcode in specs)"],
                  ["src/agent/.cache/", "Cloned app repos + UI element index JSON"],
                  [".cursor/rules/", "Cursor governance for generation quality"],
                  ["docs/presentations/", "Technical deck and POC closure PPT"],
              ])

    # ── 6. Deliverables ──
    doc.add_heading("6. POC Deliverables", level=1)
    add_bullet(doc, "Complete Playwright Agent repository (framework + agent pipeline)")
    add_bullet(doc, "28 generated, runnable .spec.ts scripts (24 billing toggle · 4 DFB)")
    add_bullet(doc, "Category CSV test data files with parsed runtime values")
    add_bullet(doc, "Agent CLI for interactive and batch generation")
    add_bullet(doc, "Reference spec catalog per category for clone-and-adapt generation")
    add_bullet(doc, "App source indexer cache workflow for real UI locators")
    add_bullet(doc, "Technical presentation: docs/presentations/agentic-system-slides.html")
    add_bullet(doc, "This handover document")

    doc.add_page_break()

    # ── 7. How to Use ──
    doc.add_heading("7. How to Use the Agent for Spec File Generation", level=1)

    doc.add_heading("7.1 Testcase Input Requirements", level=2)
    doc.add_paragraph(
        "The agent requires detailed test cases — not high-level summaries. This is the largest upstream effort."
    )
    add_bullet(doc, "Testmo cases must include full step-by-step actions and explicit expected outcomes per step")
    add_bullet(doc, "Manual QA typically spends 3–5 hours per test case reaching agent-ready detail")
    add_bullet(doc, "Without sufficient detail, generated specs are inaccurate and require heavy manual correction")
    add_bullet(doc, "JIRA scenarios may document the what — agent-ready step detail (the how) must be present")

    doc.add_heading("7.2 Supported Input Formats", level=2)
    add_bullet(doc, "CSV — src/agent/examples/sample-testcases.csv format (Case ID, Test Steps, Expected, Tags)")
    add_bullet(doc, "Excel (.xlsx, .xls)")
    add_bullet(doc, "JSON")
    add_bullet(doc, "Plain text (.txt)")

    doc.add_heading("7.3 Agent CLI Commands", level=2)
    add_table(doc,
              ["Command", "Description"],
              [
                  ["npm run agent:cli", "Interactive CLI — guided generation"],
                  ["npm run agent:file -- <path>", "Generate from a single CSV/Excel/JSON/TXT file"],
                  ["npm run agent:batch -- <folder>", "Batch process all supported files in a folder"],
                  ["npm run agent:generate -- \"description\"", "Generate from plain-text description"],
                  ["npm run agent:preview", "Preview generated code without writing files"],
                  ["npm run agent:analyze", "Analyze test cases only (no generation)"],
                  ["npm run agent -- --no-llm", "Rule-based generation only (skip all LLM calls)"],
                  ["npm run agent -- --help", "Show all CLI options"],
              ])

    doc.add_heading("7.4 Step-by-Step: Generate a New Spec from CSV", level=2)
    add_numbered(doc, "Prepare a detailed testcase row in CSV (or add to sample-testcases.csv)")
    add_numbered(doc, "Ensure runtime test data exists in src/data/<category>/ or let Call 1 (Data Enricher) fill gaps")
    add_numbered(doc, "Run: npm run agent:file -- src/agent/examples/sample-testcases.csv")
    add_numbered(doc, "Agent parses testcase → syncs app source → indexes UI elements → generates draft")
    add_numbered(doc, "Quality Guardian validates and auto-fixes (≤2 passes)")
    add_numbered(doc, "Coverage Ensurer checks step coverage (1 pass, up to 8 missing steps filled)")
    add_numbered(doc, "Output written to: src/tests/AIAgent/<category>/<TEST_ID>.spec.ts")
    add_numbered(doc, "Review generated spec; run: npx playwright test src/tests/AIAgent/<category>/<TEST_ID>.spec.ts")

    doc.add_heading("7.5 Step-by-Step: Interactive CLI", level=2)
    add_numbered(doc, "Run: npm run agent:cli")
    add_numbered(doc, "Follow prompts to select input source, testcase, and category")
    add_numbered(doc, "Review generation log for LLM tier used and any warnings")
    add_numbered(doc, "Open generated spec and verify step traceability (test.step titles reference CSV step numbers)")

    doc.add_heading("7.6 Running Generated Tests", level=2)
    add_code(doc, "# Run all AI Agent tests")
    add_code(doc, "npx playwright test src/tests/AIAgent/")
    add_code(doc, "# Run single spec")
    add_code(doc, "npx playwright test src/tests/AIAgent/billingtoggle/BT-74454.spec.ts")
    add_code(doc, "# Allure report")
    add_code(doc, "npm run test:allure")
    add_code(doc, "# Submit to Testmo")
    add_code(doc, "npm run submit:testmo")

    doc.add_heading("7.7 Adding Test Data for a New Case", level=2)
    doc.add_paragraph(
        "Runtime values (customer names, rates, office codes, etc.) live in src/data/<category>/*.csv. "
        "The CsvDataService maps CSV column headers to canonical testData keys via FieldRegistry. "
        "Generated specs load data via: dataConfig.getTestDataFromCsv(dataConfig.<category>Data, testcaseID)"
    )

    doc.add_page_break()

    # ── 8. Limitations ──
    doc.add_heading("8. Known Limitations & Upstream Gaps", level=1)
    doc.add_paragraph("The pipeline is built and working — effectiveness depends on input quality.")

    doc.add_heading("① Test Cases Lack Sufficient Detail", level=2)
    add_bullet(doc, "Testmo cases must include full step-by-step actions and expected outcomes")
    add_bullet(doc, "Manual QA spends 3–5 hours per test case — largest upstream effort")
    add_bullet(doc, "Without sufficient detail, spec generation is inaccurate")

    doc.add_heading("② Tribal Knowledge — Scenarios Without Executable Steps", level=2)
    add_bullet(doc, "Scenarios exist in JIRA — the what is often documented")
    add_bullet(doc, "Agent-ready step detail is frequently missing — the how is not")
    add_bullet(doc, "Precondition paths and business rules cannot be automated without executable steps")

    doc.add_heading("③ Limited Application Source Access", level=2)
    add_bullet(doc, "UI Element Index syncs BTMS and DME repos only")
    add_bullet(doc, "Broader repo sync needed for full product coverage")

    doc.add_heading("④ Execution Time — Monolithic Scripts", level=2)
    add_bullet(doc, "Large cases become single long-running scripts")
    add_bullet(doc, "Consider phased specs or shared fixtures for very large test cases")

    # ── 9. Reference ──
    doc.add_heading("9. Reference & Support", level=1)
    add_bullet(doc, "Technical deck: docs/presentations/agentic-system-slides.html")
    add_bullet(doc, "POC closure PPT: docs/presentations/POC-QA-Agent-Closure.pptx")
    add_bullet(doc, "Architecture rules: CLAUDE.md and .cursor/rules/agent-architecture.mdc")
    add_bullet(doc, "Reference specs: src/tests/AIAgent/billingtoggle/BT-74454.spec.ts (billing toggle), DFB-97739.spec.ts (DFB)")

    doc.add_heading("9.1 Command Quick Reference", level=2)
    add_table(doc,
              ["Task", "Command"],
              [
                  ["Install dependencies", "npm install"],
                  ["Verify Claude CLI", "claude --version"],
                  ["Interactive agent", "npm run agent:cli"],
                  ["Generate from file", "npm run agent:file -- <path>"],
                  ["Batch generate", "npm run agent:batch -- <folder>"],
                  ["Preview only", "npm run agent:preview"],
                  ["Run AI Agent tests", "npx playwright test src/tests/AIAgent/"],
                  ["Allure report", "npm run test:allure"],
                  ["Submit Testmo", "npm run submit:testmo"],
              ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Handover Document —")
    r.italic = True
    r.font.color.rgb = NAVY

    return doc


def main() -> None:
    HANDOVER_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Created: {OUT_PATH}")


if __name__ == "__main__":
    main()
