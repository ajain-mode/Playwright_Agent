"""Generate Playwright AI Agent executive report as Word document."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
    doc.add_paragraph()


def build_report(output_path: Path) -> None:
    doc = Document()
    set_document_defaults(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Playwright AI Agent\nExecutive Report")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"SunTeck TMS Quality Engineering\nPrepared: {date.today().strftime('%B %d, %Y')}"
    )
    sub_run.font.size = Pt(12)
    sub_run.italic = True

    doc.add_paragraph()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "This report describes the client transport-management application landscape, "
        "the quality-assurance challenges of automating tests against that landscape, "
        "the risks those challenges create, and how the Playwright AI Agent framework "
        "addresses them. The agent is a rule-based test-generation pipeline built on "
        "Microsoft Playwright. It converts manual test cases (CSV, JSON, plain text, or Excel) "
        "into executable TypeScript specifications backed by a mature Page Object Model (POM), "
        "with guardrails that enforce project conventions and traceability to application source code."
    )

    # 2. Client Application System Summary
    add_heading(doc, "2. Summary of the Client Application System", 1)
    doc.add_paragraph(
        "The client environment is Mode Global’s SunTeck Transport Management System (TMS) "
        "and related satellite applications used to manage freight operations end to end. "
        "Automation in this repository targets production-like workflows across multiple "
        "surfaces—not a single monolithic UI."
    )

    add_heading(doc, "2.1 Core Applications", 2)
    add_table(
        doc,
        ["Application", "Role", "Technology (indexed source)"],
        [
            [
                "BTMS (SunTeck TMS)",
                "Primary TMS: offices, customers, loads, carriers, finance, admin, commissions, EDI",
                "PHP (modetrans/mono.git — btms/php/src)",
            ],
            [
                "DME (Digital Matching Engine)",
                "Digital load matching and carrier engagement",
                "Twig templates (modetrans/dme.git)",
            ],
            [
                "TNX / TNX Rep",
                "Tender exchange: carrier tenders, execution, representative workflows",
                "Integrated via dedicated login and landing page objects",
            ],
        ],
    )

    add_heading(doc, "2.2 Functional Domains Exercised by Automation", 2)
    domains = [
        "Load lifecycle: create, edit, duplicate, view (multi-tab: Load, Pick, Drop, Carrier, Customer, EDI, Rail)",
        "Billing and finance: accounts payable, billing adjustments, payable toggles, waiting-on states",
        "Administration: offices, agents, post-automation rules, user switching",
        "Customers and carriers: search, view, edit master data",
        "EDI, commissions, sales leads, bulk changes, DAT, and API-backed scenarios",
        "DFB (Digital Freight Broker) flows spanning BTMS, TNX, and DME in a single test session",
    ]
    for d in domains:
        add_bullet(doc, d)

    add_heading(doc, "2.3 Technical Characteristics Relevant to Testing", 2)
    tech = [
        "Large, form-heavy UIs with distinct edit mode (form fields with #id) versus view mode (label/value table rows)",
        "Multi-tab detail pages where the main navigation header is hidden until the user returns to the dashboard shell",
        "Role-based access (e.g., billing-toggle user, bulk-change manager) requiring explicit account switching after login",
        "Alerts, confirmations, and finance messages that must be asserted with stable constants—not ad hoc strings",
        "Cross-application tests that open multiple browser tabs and require coordinated context switching (MultiAppManager)",
    ]
    for t in tech:
        add_bullet(doc, t)

    add_heading(doc, "2.4 Test Data and Reporting Ecosystem", 2)
    doc.add_paragraph(
        "Runtime test data is organized per category in CSV files under src/data/<category>/. "
        "Canonical manual step prose and expected results are cataloged in "
        "src/agent/examples/sample-testcases.csv for traceability. Executed tests report through "
        "Allure, JUnit XML, Playwright HTML reports, and optional Testmo submission—supporting "
        "CI visibility and test-management integration."
    )

    # 3. Problem Statement
    add_heading(doc, "3. Problem the Playwright Agent Solves", 1)
    doc.add_paragraph(
        "Manual regression of SunTeck TMS is slow, expensive, and difficult to scale. "
        "Converting written test cases into reliable Playwright automation traditionally requires "
        "senior QA engineers who understand both the business flows and a large, evolving POM library. "
        "The Playwright AI Agent closes that gap by industrializing translation from natural-language "
        "or structured test inputs into maintainable .spec.ts files."
    )

    add_heading(doc, "3.1 Specific Pain Points", 2)
    pains = [
        "High authoring cost: Each new case may need dozens of navigation steps, tab switches, data entry, and assertions aligned to written expected results.",
        "Locator fragility: Guessing XPath or partial IDs leads to flaky tests; BTMS and DME UIs change with application releases.",
        "Convention drift: Without enforcement, generated code may hardcode strings, bypass POMs, omit navigateToBaseUrl before header actions, or skip mandatory user switches (e.g., billing toggle).",
        "Context complexity: Steps must know whether the UI is in edit or view mode, which tab is active, and which application (BTMS vs TNX vs DME) owns the action.",
        "Traceability gap: Manual cases in spreadsheets often diverge from automated specs over time, making audits and defect reproduction harder.",
        "POM sprawl: New UI fields require new page-object methods; duplicating logic across specs increases maintenance.",
    ]
    for p in pains:
        add_bullet(doc, p)

    # 4. Key Risks
    add_heading(doc, "4. Key Risks Inherent in the Problem", 1)
    doc.add_paragraph(
        "If test automation is built manually without the controls this agent provides, "
        "the organization faces the following material risks:"
    )

    risk_rows = [
        [
            "Flaky / brittle automation",
            "Incorrect or invented locators break on minor UI changes; tests fail intermittently and erode trust in CI.",
        ],
        [
            "False confidence",
            "Specs that omit explicit expected validations (console.log only) pass while business rules regress.",
        ],
        [
            "Security & compliance exposure",
            "Hardcoded credentials or environment secrets committed in generated code.",
        ],
        [
            "Regression in critical financial flows",
            "Billing toggle, waiting-on, and payable status mistakes can miss revenue or carrier-payment defects.",
        ],
        [
            "Multi-app workflow breakage",
            "DFB scenarios that span BTMS, TNX, and DME fail when tab/context switching is wrong.",
        ],
        [
            "Human POM corruption",
            "Automated tools overwriting hand-maintained page objects destroy team IP and destabilize the whole suite.",
        ],
        [
            "Maintenance debt",
            "Inconsistent patterns (locators in specs, magic numbers, duplicate navigation) multiply fix cost per release.",
        ],
        [
            "Slow time-to-coverage",
            "Backlog of manual cases grows faster than automation capacity; release risk accumulates.",
        ],
    ]
    add_table(doc, ["Risk", "Impact"], risk_rows)

    # 5. Solution
    add_heading(doc, "5. Solution Provided by the Playwright AI Agent", 1)
    doc.add_paragraph(
        "The agent is a three-stage, rule-based pipeline (LLM optional for clone-and-adapt edge cases) "
        "that generates production-ready Playwright tests under src/tests/AIAgent/<category>/."
    )

    add_heading(doc, "5.1 Generation Pipeline", 2)
    pipeline = [
        "TestCaseParser — Parses JSON, CSV, XLSX, or plain text; detects category and extracts field values via FieldRegistry.",
        "RepoCloneManager + AppSourceIndexer — Clones BTMS (PHP) and DME (Twig) repos; indexes real HTML element IDs; never fabricates locators when source is available.",
        "CsvDataService — Ensures category CSV rows exist for runtime testData.* consumption.",
        "CodeGenerator — High-match reference cloning (≥70%) or step-by-step assembly via StepProcessor, StepMappings (130+ patterns), and POMMethodMatcher.",
        "SpecValidator — Sanitizer pre-pass (13 rules) plus guardrails (NAV-001, CAT-BT-001, no locators in specs, no hardcoded assertions, etc.) with auto-fix where safe.",
        "PageObjectWriter — Adds only @author AI Agent methods; protects human-authored POM code.",
    ]
    for i, step in enumerate(pipeline, 1):
        add_numbered(doc, step)

    add_heading(doc, "5.2 How the Solution Mitigates Each Risk Area", 2)
    mitigations = [
        ("Locator accuracy", "AppSourceIndexer scores elements by stability; POMMethodMatcher prefers #id from application source."),
        ("Assertion quality", "Every explicit expected from the testcase maps to expect/expect.soft or structured alert handling; constants from globalConstants.ts and testData from CSV."),
        ("Navigation reliability", "Mandatory navigateToBaseUrl() before header navigation from detail/form pages (NAV-001 auto-fix)."),
        ("Category rules", "e.g., billingtoggle specs auto-inject USER_ROLES.BILLINGTOGGLE_USER switch (CAT-BT-001)."),
        ("View vs edit awareness", "StepProcessor tracks isEditMode and currentTab to choose correct locator strategy."),
        ("Traceability", "sample-testcases.csv + nested test.step titles aligned to CSV step numbers."),
        ("Safe POM evolution", "Human-authored methods are read-only; agent may only add or modify AI-tagged methods."),
        ("Scale", "CLI modes: agent:file, agent:batch, agent:preview, agent:generate for repeatable bulk generation."),
    ]
    for title, detail in mitigations:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(detail)

    add_heading(doc, "5.3 Supported Test Categories", 2)
    doc.add_paragraph(
        "dfb, edi, commission, salesLead, banyan, carrier, bulkChange, dat, "
        "nonOperationalLoads, api, billingtoggle — each with dedicated CSV data and reference specs."
    )

    add_heading(doc, "5.4 Operational Outcomes", 2)
    outcomes = [
        "Faster conversion of manual cases to executable Playwright specs with consistent structure.",
        "Reduced locator guesswork by grounding automation in cloned application source.",
        "Enforced engineering standards via SpecValidator guardrails and sanitizer rules.",
        "Integrated reporting (Allure, JUnit, Testmo) for pipeline and test-management visibility.",
        "Extensible POM layer shared across human-written and agent-generated tests.",
    ]
    for o in outcomes:
        add_bullet(doc, o)

    # 6. Conclusion
    add_heading(doc, "6. Conclusion", 1)
    doc.add_paragraph(
        "The client TMS ecosystem is complex, multi-application, and business-critical. "
        "Traditional manual test automation does not scale safely without strong conventions "
        "and source-grounded locators. The Playwright AI Agent addresses that gap with a "
        "disciplined, rule-based generation pipeline, comprehensive guardrails, and deep "
        "integration with existing Page Objects and test data—turning written test cases into "
        "maintainable, reviewable Playwright automation while controlling the key risks of "
        "flakiness, convention drift, and maintenance debt."
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "Document generated from Playwright_Agent repository context. "
        "For technical details, see CLAUDE.md and .cursor/rules/agent-architecture.mdc."
    ).italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    out = root / "docs" / "Playwright_AI_Agent_Executive_Report.docx"
    build_report(out)
